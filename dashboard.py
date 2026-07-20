import streamlit as st
import pandas as pd
import folium
from streamlit_folium import st_folium
import plotly.express as px
from io import BytesIO

from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer
)

from reportlab.lib.styles import getSampleStyleSheet

from docx import Document

def generate_pdf(df):

    buffer = BytesIO()

    doc = SimpleDocTemplate(buffer)

    styles = getSampleStyleSheet()

    story = []

    story.append(
        Paragraph(
            "<b>Urban Decision Intelligence Platform</b>",
            styles["Title"]
        )
    )

    story.append(
        Paragraph(
            "Executive Decision Report",
            styles["Heading2"]
        )
    )

    story.append(Spacer(1,20))

    story.append(
        Paragraph(
            f"<b>Total Parking Facilities :</b> {len(df)}",
            styles["BodyText"]
        )
    )

    story.append(
        Paragraph(
            f"<b>Critical Locations :</b> {(df['Priority']=='Critical').sum()}",
            styles["BodyText"]
        )
    )

    story.append(
        Paragraph(
            f"<b>Average Systemic Stress :</b> {df['SystemicStress'].mean():.2f}",
            styles["BodyText"]
        )
    )

    story.append(
        Paragraph(
            f"<b>Average Vehicle Count :</b> {df['VehicleCount'].mean():,.0f}",
            styles["BodyText"]
        )
    )

    story.append(Spacer(1,20))

    story.append(
        Paragraph(
            "<b>Executive Summary</b>",
            styles["Heading2"]
        )
    )

    story.append(
        Paragraph(
            """
The Urban Decision Intelligence Platform integrates
parking and traffic information to identify high-stress
parking facilities and provide prescriptive recommendations
for urban planners and government agencies.

Facilities with Critical priority require immediate
operational attention to improve traffic flow and
parking efficiency.
""",
            styles["BodyText"]
        )
    )

    doc.build(story)

    buffer.seek(0)

    return buffer
def generate_docx(df):

    doc = Document()

    doc.add_heading(
        "Urban Decision Intelligence Platform",
        0
    )

    doc.add_heading(
        "Executive Decision Report",
        level=1
    )

    doc.add_paragraph(
        f"Total Parking Facilities : {len(df)}"
    )

    doc.add_paragraph(
        f"Critical Locations : {(df['Priority']=='Critical').sum()}"
    )

    doc.add_paragraph(
        f"Average Stress : {df['SystemicStress'].mean():.2f}"
    )

    doc.add_paragraph(
        f"Average Vehicle Count : {df['VehicleCount'].mean():,.0f}"
    )

    doc.add_heading(
        "Executive Summary",
        level=2
    )

    doc.add_paragraph(
        """
The Urban Decision Intelligence Platform analyzes
traffic intensity and parking conditions to support
government decision making through systemic stress
evaluation and prescriptive recommendations.
"""
    )

    buffer = BytesIO()

    doc.save(buffer)

    buffer.seek(0)

    return buffer



# ======================================
# Page Configuration
# ======================================
st.set_page_config(
    page_title="Government Decision Intelligence Dashboard",
    page_icon="🏛️",
    layout="wide"
)

# ======================================
# Load Data
# ======================================
df = pd.read_csv("output/final_decision_intelligence.csv")

# ======================================
# Sidebar Navigation
# ======================================
st.sidebar.title("🏛️ UDIP")
st.sidebar.caption("Urban Decision Intelligence Platform")

st.sidebar.markdown("---")

st.sidebar.subheader("Navigation")

selected = st.sidebar.radio(
    "",
    [
        "🏠 Dashboard",
        "📊 Parking Analysis",
        "🗺️ Traffic Intelligence",
        "🚦 Decision Engine",
        "📑 Reports"
    ]
)

st.sidebar.markdown("---")

st.sidebar.subheader("System Status")
st.sidebar.success("🟢 Decision Engine Online")
st.sidebar.info("Dataset Loaded")
st.sidebar.write(f"Parking Facilities : **{len(df)}**")

st.sidebar.markdown("---")
st.sidebar.caption("Government Decision Support System")

# ======================================
# Dashboard
# ======================================
if selected == "🏠 Dashboard":

    st.title("🏛️ Government Decision Intelligence Dashboard")
    st.markdown("### Decision Support System for Urban Parking Management")

    st.divider()

    # ======================================
    # KPI Calculations
    # ======================================
    total = len(df)
    critical = (df["Priority"] == "Critical").sum()
    high = (df["Priority"] == "High").sum()
    moderate = (df["Priority"] == "Moderate").sum()
    low = (df["Priority"] == "Low").sum()

    # ======================================
    # KPI Cards
    # ======================================
    col1, col2, col3, col4, col5 = st.columns(5)

    col1.metric("Parking Facilities", total)
    col2.metric("🔴 Critical", critical)
    col3.metric("🟠 High", high)
    col4.metric("🟡 Moderate", moderate)
    col5.metric("🟢 Low", low)

    st.divider()

    # ======================================
    # Top Priority Parking
    # ======================================
    st.subheader("🏆 Top Priority Parking Locations")

    priority_order = {
        "Critical": 4,
        "High": 3,
        "Moderate": 2,
        "Low": 1
    }

    temp = df.copy()
    temp["PriorityRank"] = temp["Priority"].map(priority_order)

    top5 = (
        temp.sort_values(
            by=["PriorityRank", "SystemicStress"],
            ascending=[False, False]
        )
        .head(5)
    )

    st.dataframe(
        top5[
            [
                "ParkingName",
                "Priority",
                "SystemicStress",
                "ImmediateActions"
            ]
        ],
        use_container_width=True,
        hide_index=True
    )

    st.divider()

    # ======================================
    # Interactive Map
    # ======================================
    st.subheader("🗺️ Urban Parking Intelligence Map")

    m = folium.Map(
        location=[52.4862, -1.8904],
        zoom_start=13,
        tiles="CartoDB positron"
    )

    color_map = {
        "Critical": "red",
        "High": "orange",
        "Moderate": "blue",
        "Low": "green"
    }

    for _, row in df.iterrows():

        popup = f"""
        <h4>{row['ParkingName']}</h4>

        <b>Priority:</b> {row['Priority']}<br>

        <b>Systemic Stress:</b> {row['SystemicStress']:.2f}<br><br>

        <b>Reason</b><br>
        {row['Reason']}<br><br>

        <b>Immediate Actions</b><br>
        {row['ImmediateActions']}<br><br>

        <b>Long-Term Considerations</b><br>
        {row['LongTermConsiderations']}
        """

        folium.Marker(
            location=[row["Latitude"], row["Longitude"]],
            tooltip=row["ParkingName"],
            popup=folium.Popup(popup, max_width=400),
            icon=folium.Icon(
                color=color_map.get(row["Priority"], "gray"),
                icon="car",
                prefix="fa"
            )
        ).add_to(m)

    st_folium(
        m,
        use_container_width=True,
        height=600
    )

    st.divider()

    # ======================================
    # Parking Facilities Table
    # ======================================
    st.subheader("📋 Parking Facilities")

    st.dataframe(
        df[
            [
                "ParkingName",
                "Priority",
                "SystemicStress",
                "TrafficLevel",
                "Reason"
            ]
        ],
        use_container_width=True,
        hide_index=True
    )

# ======================================
# Parking Analysis
# ======================================
elif selected == "📊 Parking Analysis":

    st.title("📊 Parking Analysis Dashboard")

    # ==========================
    # Executive KPIs
    # ==========================

    col1, col2, col3 = st.columns(3)

    col1.metric(
        "Average Systemic Stress",
        f"{df['SystemicStress'].mean():.2f}"
    )

    col2.metric(
        "Highest Stress",
        f"{df['SystemicStress'].max():.2f}"
    )

    col3.metric(
        "Average Vehicle Count",
        f"{df['VehicleCount'].mean():,.0f}"
    )

    st.divider()

    # ==========================
    # Charts
    # ==========================

    left, right = st.columns(2)

    with left:

        priority = (
            df["Priority"]
            .value_counts()
            .reset_index()
        )

        priority.columns = ["Priority", "Count"]

        fig = px.bar(
            priority,
            x="Priority",
            y="Count",
            color="Priority",
            title="Priority Distribution"
        )

        st.plotly_chart(fig, use_container_width=True)

    with right:

        traffic = (
            df["TrafficLevel"]
            .value_counts()
            .reset_index()
        )

        traffic.columns = ["Traffic Level", "Count"]

        fig = px.pie(
            traffic,
            names="Traffic Level",
            values="Count",
            title="Traffic Level Distribution"
        )

        st.plotly_chart(fig, use_container_width=True)

    st.divider()

    fig = px.histogram(
        df,
        x="SystemicStress",
        nbins=10,
        title="Systemic Stress Distribution"
    )

    st.plotly_chart(fig, use_container_width=True)

    st.divider()

    st.subheader("🏆 Top 10 Highest Stress Parking Facilities")

    st.dataframe(
        df.sort_values(
            "SystemicStress",
            ascending=False
        ).head(10)[
            [
                "ParkingName",
                "Priority",
                "SystemicStress",
                "TrafficLevel"
            ]
        ],
        use_container_width=True,
        hide_index=True
    )

# ======================================
# Traffic Intelligence
# ======================================
elif selected == "🗺️ Traffic Intelligence":

    st.title("🗺️ Traffic Intelligence Dashboard")

    st.markdown("### Urban Traffic Monitoring & Congestion Analysis")

    st.divider()

    # ======================================
    # Executive KPIs
    # ======================================

    col1, col2, col3 = st.columns(3)

    col1.metric(
        "Average Vehicle Count",
        f"{df['VehicleCount'].mean():,.0f}"
    )

    col2.metric(
        "Maximum Vehicle Count",
        f"{df['VehicleCount'].max():,.0f}"
    )

    col3.metric(
        "Very High Traffic Areas",
        (df["TrafficLevel"] == "Very High").sum()
    )

    st.divider()

    # ======================================
    # Traffic Level Distribution
    # ======================================

    traffic = (
        df["TrafficLevel"]
        .value_counts()
        .reset_index()
    )

    traffic.columns = ["Traffic Level", "Count"]

    fig = px.bar(
        traffic,
        x="Traffic Level",
        y="Count",
        color="Traffic Level",
        title="Traffic Level Distribution"
    )

    st.plotly_chart(
        fig,
        use_container_width=True
    )

    st.divider()

    # ======================================
    # Vehicle Count Heat Map
    # ======================================

    st.subheader("🗺️ Vehicle Count Map")

    traffic_map = folium.Map(
        location=[52.4862, -1.8904],
        zoom_start=13,
        tiles="CartoDB positron"
    )

    color_map = {
        "Very High": "red",
        "High": "orange",
        "Medium": "blue",
        "Low": "green"
    }

    for _, row in df.iterrows():

        folium.CircleMarker(
            location=[row["Latitude"], row["Longitude"]],
            radius=max(6, row["VehicleCount"] / 2500),
            color=color_map.get(row["TrafficLevel"], "gray"),
            fill=True,
            fill_opacity=0.7,
            popup=f"""
<b>{row['ParkingName']}</b><br>
Nearest Road: {row['NearestRoad']}<br>
Vehicle Count: {row['VehicleCount']:,}<br>
Traffic Level: {row['TrafficLevel']}
"""
        ).add_to(traffic_map)

    st_folium(
        traffic_map,
        use_container_width=True,
        height=600
    )

    st.divider()

    # ======================================
    # Top Congested Roads
    # ======================================

    st.subheader("🚦 Top Congested Roads")

    st.dataframe(

        df.sort_values(
            "VehicleCount",
            ascending=False
        )[

            [
                "NearestRoad",
                "VehicleCount",
                "TrafficLevel",
                "ParkingName"
            ]

        ].head(10),

        hide_index=True,
        use_container_width=True
    )

# ======================================
# Decision Engine
# ======================================
elif selected == "🚦 Decision Engine":

    st.title("🚦 Government Decision Engine")

    st.markdown("### Select a parking facility to view AI-generated recommendations")

    parking = st.selectbox(
        "Parking Facility",
        sorted(df["ParkingName"].unique())
    )

    row = df[df["ParkingName"] == parking].iloc[0]

    st.divider()

    col1, col2, col3 = st.columns(3)

    col1.metric("Priority", row["Priority"])
    col2.metric("Systemic Stress", f"{row['SystemicStress']:.2f}")
    col3.metric("Traffic Level", row["TrafficLevel"])

    st.divider()

    st.subheader("🧠 Reason")

    st.info(row["Reason"])

    st.subheader("⚡ Immediate Actions")

    st.success(row["ImmediateActions"])

    st.subheader("🏗 Long-Term Considerations")

    st.warning(row["LongTermConsiderations"])

    st.divider()

    st.subheader("📊 Additional Information")

    left, right = st.columns(2)

    with left:

        st.write(f"**Peak Hour :** {row['PeakHour']}")
        st.write(f"**Peak Day :** {row['PeakDay']}")
        st.write(f"**Priority :** {row['Priority']}")

    with right:

        st.write(f"**Traffic Level :** {row['TrafficLevel']}")
        st.write(f"**Vehicle Count :** {row['VehicleCount']}")
        st.write(f"**Systemic Stress :** {row['SystemicStress']:.2f}")

# ======================================
# Reports
# ======================================
elif selected == "📑 Reports":

    st.title("📑 Executive Reports")

    st.markdown("### Urban Decision Intelligence Executive Summary")

    st.divider()

    # ==============================
    # Executive KPIs
    # ==============================

    col1, col2, col3, col4 = st.columns(4)

    col1.metric(
        "Parking Facilities",
        len(df)
    )

    col2.metric(
        "Critical Locations",
        (df["Priority"] == "Critical").sum()
    )

    col3.metric(
        "Average Stress",
        f"{df['SystemicStress'].mean():.2f}"
    )

    col4.metric(
        "Average Vehicle Count",
        f"{df['VehicleCount'].mean():,.0f}"
    )

    st.divider()

    # ==============================
    # Executive Summary
    # ==============================

    st.subheader("🏛 Executive Summary")

    summary = f"""
The analysis evaluated **{len(df)} parking facilities** by integrating
parking characteristics with surrounding traffic conditions.

The overall average Systemic Stress Index is
**{df['SystemicStress'].mean():.2f}**.

A total of **{(df['Priority']=='Critical').sum()}**
locations require immediate intervention.

The platform recommends prioritizing these facilities
through traffic management, dynamic parking policies,
and infrastructure improvements.
"""

    st.success(summary)

    st.divider()

    # ==============================
    # Top Priority Locations
    # ==============================

    st.subheader("🚨 Highest Priority Locations")

    report = df.sort_values(
        "SystemicStress",
        ascending=False
    )

    st.dataframe(

        report[
            [
                "ParkingName",
                "Priority",
                "SystemicStress",
                "TrafficLevel",
                "Reason"
            ]
        ].head(10),

        hide_index=True,
        use_container_width=True
    )

    st.divider()

    # ==============================
    # Download Report
    # ==============================

    csv = df.to_csv(index=False).encode("utf-8")

    st.download_button(
        "⬇ Download Decision Intelligence Report",
        csv,
        "Decision_Intelligence_Report.csv",
        "text/csv"
    )

    pdf = generate_pdf(df)

st.download_button(
    "📄 Download Executive Report (PDF)",
    pdf,
    "Executive_Report.pdf",
    mime="application/pdf"
)

docx = generate_docx(df)

st.download_button(
    "📝 Download Executive Report (DOCX)",
    docx,
    "Executive_Report.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)